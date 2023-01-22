from docx import Document


class WordDoc:

   
    def __init__(self, templateName):
        self.document = Document(templateName)
    # code to display table inside word file
    # Bring in user created template file, new data tables will be appended to the file and saved in a new location
    def Table(self, tableHeader, tableBody):

        table = self.document.add_table(
            rows=1, cols=len(tableHeader))

        # Create table header
        hdr_cells = table.rows[0].cells
        for count, string in enumerate(tableHeader):
            hdr_cells[count].text = tableHeader[count]

        # Create table body

        for row in tableBody:
            row_cells = table.add_row().cells
            for count, string in enumerate(row):
                row_cells[count].text = string

    # Progress this month, Started and completed activity lists
    def actSC(self, started_activities: tuple, completed_activities: tuple):

        self.document.add_heading('Progress Made This Month', level=1)

        # Started Activities
        self.document.add_heading('Activities Started This Month', level=2)
        for string in started_activities:
            self.document.add_paragraph(string, style='List Bullet')

        # Completed Activities
        self.document.add_heading('Activities Completed This Month', level=2)
        for string in completed_activities:
            self.document.add_paragraph(string, style='List Bullet')

    # Critical path activities for the next 30 days
    def next30CP(self, critical30day):
        self.document.add_heading(
            'Activities/Progress/Events Anticipated to Affect the Critical Path in the Next 30 Days', level=1)
        self.document.add_paragraph(
            'Progress of the following activities at or above the scheduled rate is essential to avoid float degradation.')
        for string in critical30day:
            self.document.add_paragraph(string, style='List Bullet')

    # Display changed activity name table
    def table_ActivityName(self, header, body):
        self.document.add_heading("Activity Name", level=2)
        self.Table(header, body)

    # Display added activity table
    def table_AddedActivities(self, header, body):
        self.document.add_heading("Added Activities", level=2)
        self.Table(header, body)

    # Display deleted activity table
    def table_DeletedActivities(self, header, body):
        self.document.add_heading("Deleted Activities", level=2)
        self.Table(header, body)

    # Display changed original duration table
    def table_ODur(self, header, body):
        self.document.add_heading("Original Duration", level=2)
        self.Table(header, body)

    # Display added successor relationship table
    def table_AddedSucc(self, header, body):
        self.document.add_heading(
            "Added Successor Activity Relationships", level=2)
        self.Table(header, body)

    # Display deleted successor table
    def table_DeletedSucc(self, header, body):
        self.document.add_heading(
            "Deleted Successor Activity Relationships", level=2)
        self.Table(header, body)

    # Display all activity and relationship modifications
    # Arguments in form [[Header],[Body]]
    def activityMods(self, actNameTable, addedActTable, deletedActTable, addedSuccTable, deletedSuccTable, changedOdurTable):
        self.document.add_heading("Activity and Relationship Modifications", level=1)
        self.table_ActivityName(actNameTable[0], actNameTable[1])
        self.table_AddedActivities(addedActTable[0], addedActTable[1])
        self.table_DeletedActivities(deletedActTable[0], deletedActTable[1])
        self.table_AddedSucc(addedSuccTable[0], addedSuccTable[1])
        self.table_DeletedSucc(deletedSuccTable[0], deletedSuccTable[1])
        self.table_ODur(changedOdurTable[0], changedOdurTable[1])

    def saveFile(self, saveName):
        # Save document
        self.document.save(saveName)
